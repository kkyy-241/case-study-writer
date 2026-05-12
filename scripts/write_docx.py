"""Render a UTF-8 text or Markdown-like draft to DOCX."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path


DEFAULT_FONT = "Microsoft YaHei"
HEADING_FONT = "Microsoft YaHei"
IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")
TEMPLATE_FONT = "SimSun"
TEMPLATE_HEADING_FONT = "SimSun"
TEMPLATE_TITLE_FONT = "YouYuan"
HEADER_LEFT_TEXT = "商学院教学案例库"


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def parse_line(line: str) -> tuple[str, str]:
    stripped = line.strip()
    heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
    if heading:
        level = str(min(len(heading.group(1)), 3))
        return f"Heading {level}", clean_inline_markdown(heading.group(2))
    return "Normal", clean_inline_markdown(stripped)


def set_east_asia_font(style, font_name: str) -> None:
    from docx.oxml.ns import qn

    style.font.name = font_name
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def infer_company_suffix(*paths: Path) -> str:
    for path in paths:
        stem = path.stem
        for prefix in ("案例正文_", "案例使用说明_"):
            if stem.startswith(prefix):
                suffix = stem[len(prefix) :].strip()
                if suffix:
                    return suffix
    return ""


def set_run_east_asia_font(run, font_name: str) -> None:
    from docx.oxml.ns import qn

    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)


def set_cell_margins(cell, top: int = 0, start: int = 0, bottom: int = 0, end: int = 0) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "nil")


def set_table_indent_zero(table) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")


def add_header_run(paragraph, text: str):
    from docx.shared import Pt

    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    set_run_east_asia_font(run, TEMPLATE_FONT)
    return run


def configure_header(section, company_suffix: str) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    header = section.header
    for paragraph in header.paragraphs:
        paragraph.text = ""

    table = header.add_table(rows=1, cols=2, width=Cm(14.66))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    remove_table_borders(table)
    set_table_indent_zero(table)

    left_cell, right_cell = table.rows[0].cells
    left_cell.width = Cm(7.33)
    right_cell.width = Cm(7.33)

    for cell in (left_cell, right_cell):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_header_run(left_paragraph, HEADER_LEFT_TEXT)

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_header_run(right_paragraph, company_suffix)


def configure_template_document(document, company_suffix: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.65)
    section.bottom_margin = Cm(2.15)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.35)
    section.footer_distance = Cm(1.55)

    styles = document.styles
    normal = styles["Normal"]
    set_east_asia_font(normal, TEMPLATE_FONT)
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(24)

    for level, size in ((1, 14), (2, 12), (3, 12)):
        style = styles[f"Heading {level}"]
        set_east_asia_font(style, TEMPLATE_HEADING_FONT)
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.line_spacing = Pt(20)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Pt(0)

    configure_header(section, company_suffix)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)


def infer_document_kind(source: Path, destination: Path, text: str) -> str:
    joined = f"{source.name} {destination.name} {text[:80]}"
    if "使用说明" in joined or "teaching" in joined.lower() or "note" in joined.lower():
        return "案例使用说明"
    return "案例正文"


def add_template_title(document, title: str, kind: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    label = document.add_paragraph()
    label.paragraph_format.first_line_indent = Pt(0)
    label.paragraph_format.space_after = Pt(14)
    label_run = label.add_run(f"{kind}：")
    label_run.bold = False
    label_run.font.name = "Times New Roman"
    label_run.font.size = Pt(15)
    set_run_east_asia_font(label_run, TEMPLATE_TITLE_FONT)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_after = Pt(30)
    run = paragraph.add_run(title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    set_run_east_asia_font(run, TEMPLATE_TITLE_FONT)


def write_docx(source: Path, destination: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Run scripts/check_dependencies.py.") from exc

    text = source.read_text(encoding="utf-8-sig").replace("\r\n", "\n").replace("\r", "\n")
    document = Document()
    company_suffix = infer_company_suffix(destination, source)
    configure_template_document(document, company_suffix)
    kind = infer_document_kind(source, destination, text)
    title_written = False

    for raw_line in text.split("\n"):
        if not raw_line.strip():
            continue

        image_match = IMAGE_PATTERN.match(raw_line.strip())
        if image_match:
            image_path = Path(image_match.group("path"))
            if not image_path.is_absolute():
                image_path = (source.parent / image_path).resolve()
            document.add_picture(str(image_path), width=Cm(14.5))
            alt_text = clean_inline_markdown(image_match.group("alt"))
            if alt_text:
                caption = document.add_paragraph(style="Normal")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.first_line_indent = None
                caption.add_run(alt_text)
            continue

        style, content = parse_line(raw_line)
        if not content:
            continue
        if not title_written and style == "Heading 1":
            add_template_title(document, content, kind)
            title_written = True
            continue

        paragraph = document.add_paragraph(style=style)
        if style != "Normal":
            paragraph.paragraph_format.first_line_indent = None
        paragraph.add_run(content)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a UTF-8 text or Markdown-like draft to DOCX.")
    parser.add_argument("source", type=Path, help="Source .txt or .md file.")
    parser.add_argument("destination", type=Path, help="Destination .docx file.")
    args = parser.parse_args()

    write_docx(args.source, args.destination)
    print(f"Wrote {args.destination}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
